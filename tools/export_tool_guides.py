"""
Tool Guide Export Tool
=======================
Compiles all 10 tool guides into a single Word document and PDF.
Each tool guide starts on a fresh page.

Usage:
    python tools/export_tool_guides.py              # Generate both Word and PDF
    python tools/export_tool_guides.py --word-only   # Word only
    python tools/export_tool_guides.py --pdf-only    # PDF only (requires Word)

Output:
    output/Strategy-Tool-Guides-Complete.docx
    output/Strategy-Tool-Guides-Complete.pdf
"""

import os
import re
import sys
import subprocess
import argparse

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if PROJECT_ROOT not in sys.path:
    sys.path.insert(0, PROJECT_ROOT)

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from PIL import Image

# ── Paths ───────────────────────────────────────────────────────────────

GUIDES_DIR = os.path.join(PROJECT_ROOT, 'output', 'tool-guides')
GRAPHICS_DIR = os.path.join(PROJECT_ROOT, 'output', 'graphics', 'tool-guides')
OUTPUT_DOCX = os.path.join(PROJECT_ROOT, 'output', 'Strategy-Tool-Guides-Complete.docx')
OUTPUT_PDF = os.path.join(PROJECT_ROOT, 'output', 'Strategy-Tool-Guides-Complete.pdf')

# Guide ordering and figure mapping
# (filename, figure_png, figure_number, figure_title, alt_text)
GUIDE_ORDER = [
    {
        'file': 'EFE-Matrix.md',
        'figure': 'fig-tg-2-efe-example.png',
        'figure_number': 'TG.2',
        'figure_title': 'EFE Matrix Construction Example',
        'alt_text': 'Table showing completed EFE Matrix example with opportunities, threats, weights, ratings, and weighted scores',
    },
    {
        'file': 'IFE-Matrix.md',
        'figure': 'fig-tg-3-ife-example.png',
        'figure_number': 'TG.3',
        'figure_title': 'IFE Matrix Construction Example',
        'alt_text': 'Table showing completed IFE Matrix example with strengths, weaknesses, weights, ratings, and weighted scores',
    },
    {
        'file': 'CPM.md',
        'figure': 'fig-tg-4-cpm-example.png',
        'figure_number': 'TG.4',
        'figure_title': 'Competitive Profile Matrix Example',
        'alt_text': 'Table showing CPM example comparing three competitors across critical success factors',
    },
    {
        'file': 'SWOT-Matrix.md',
        'figure': 'fig-tg-5-swot-example.png',
        'figure_number': 'TG.5',
        'figure_title': 'SWOT Matrix Strategy Development Example',
        'alt_text': 'Completed SWOT Matrix showing SO, WO, ST, and WT strategies developed from SWOT factors',
    },
    {
        'file': 'SPACE-Matrix.md',
        'figure': 'fig-tg-6-space-example.png',
        'figure_number': 'TG.6',
        'figure_title': 'SPACE Matrix Plotting Example',
        'alt_text': 'SPACE Matrix example with directional vector plotted from dimension scores',
    },
    {
        'file': 'BCG-Matrix.md',
        'figure': 'fig-tg-1-bcg-example.png',
        'figure_number': 'TG.1',
        'figure_title': 'BCG Matrix Worked Example',
        'alt_text': 'BCG Matrix example showing four product divisions plotted in Stars, Question Marks, Cash Cows, and Dogs quadrants',
    },
    {
        'file': 'IE-Matrix.md',
        'figure': 'fig-tg-7-ie-example.png',
        'figure_number': 'TG.7',
        'figure_title': 'IE Matrix Positioning Example',
        'alt_text': 'IE Matrix example showing division positioning based on IFE and EFE total weighted scores',
    },
    {
        'file': 'Grand-Strategy-Matrix.md',
        'figure': 'fig-tg-8-grand-strategy-example.png',
        'figure_number': 'TG.8',
        'figure_title': 'Grand Strategy Matrix Application Example',
        'alt_text': 'Grand Strategy Matrix example showing company positioning based on competitive position and market growth',
    },
    {
        'file': 'QSPM.md',
        'figure': 'fig-tg-9-qspm-example.png',
        'figure_number': 'TG.9',
        'figure_title': 'QSPM Decision Analysis Example',
        'alt_text': 'QSPM example showing attractiveness scores for two strategy alternatives across key factors',
    },
    {
        'file': 'Perceptual-Map.md',
        'figure': 'fig-tg-10-perceptual-map-example.png',
        'figure_number': 'TG.10',
        'figure_title': 'Perceptual Map Construction Example',
        'alt_text': 'Perceptual map example showing competitor positioning on two strategic dimensions',
    },
]

# ── Design Constants ────────────────────────────────────────────────────

NAVY = RGBColor(0x1B, 0x2A, 0x4A)
STEEL = RGBColor(0x3D, 0x5A, 0x80)
TEXT_COLOR = RGBColor(0x21, 0x25, 0x29)
TEXT_SECONDARY = RGBColor(0x49, 0x50, 0x57)
FONT_FAMILY = 'Segoe UI'
FONT_BODY_SIZE = Pt(11)


# ── Document Builder ────────────────────────────────────────────────────

def create_document():
    """Create a styled Word document."""
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_FAMILY
    font.size = FONT_BODY_SIZE
    font.color.rgb = TEXT_COLOR
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15

    for level, (size, color, bold) in enumerate([
        (Pt(22), NAVY, True),
        (Pt(16), NAVY, True),
        (Pt(13), STEEL, True),
        (Pt(11.5), STEEL, True),
    ], start=1):
        hstyle = doc.styles[f'Heading {level}']
        hfont = hstyle.font
        hfont.name = FONT_FAMILY
        hfont.size = size
        hfont.color.rgb = color
        hfont.bold = bold
        hpf = hstyle.paragraph_format
        hpf.space_before = Pt(18 if level <= 2 else 12)
        hpf.space_after = Pt(6)
        hpf.keep_with_next = True

    return doc


def add_title_page(doc):
    """Add a title page for tool guides collection."""
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Strategic Management Tool Guides')
    run.font.size = Pt(32)
    run.font.color.rgb = NAVY
    run.font.name = FONT_FAMILY
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Frameworks for Strategic Analysis')
    run.font.size = Pt(18)
    run.font.color.rgb = STEEL
    run.font.name = FONT_FAMILY
    run.italic = True

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('MBA Strategy Course')
    run.font.size = Pt(14)
    run.font.color.rgb = TEXT_SECONDARY
    run.font.name = FONT_FAMILY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Grand Canyon University')
    run.font.size = Pt(14)
    run.font.color.rgb = TEXT_SECONDARY
    run.font.name = FONT_FAMILY

    for _ in range(4):
        doc.add_paragraph()

    # Tool guide list
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Included Guides:')
    run.font.size = Pt(11)
    run.font.color.rgb = TEXT_SECONDARY
    run.font.name = FONT_FAMILY
    run.bold = True

    guide_names = [
        'EFE Matrix', 'IFE Matrix', 'Competitive Profile Matrix',
        'SWOT Matrix', 'SPACE Matrix', 'BCG Growth-Share Matrix',
        'IE Matrix', 'Grand Strategy Matrix', 'QSPM',
        'Perceptual Map',
    ]
    for name in guide_names:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(name)
        run.font.size = Pt(10)
        run.font.color.rgb = TEXT_SECONDARY
        run.font.name = FONT_FAMILY

    doc.add_page_break()


def add_toc(doc):
    """Add a Table of Contents."""
    doc.add_heading('Table of Contents', level=1)

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar1 = run._r.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    run._r.append(fldChar1)

    run2 = paragraph.add_run()
    instrText = run2._r.makeelement(qn('w:instrText'), {})
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run2._r.append(instrText)

    run3 = paragraph.add_run()
    fldChar2 = run3._r.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'separate'})
    run3._r.append(fldChar2)

    run4 = paragraph.add_run()
    run4._r.text = 'Press F9 to update table of contents'
    run4.font.color.rgb = TEXT_SECONDARY

    run5 = paragraph.add_run()
    fldChar3 = run5._r.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    run5._r.append(fldChar3)

    doc.add_page_break()


def add_image_to_doc(doc, img_path, alt_text=''):
    """Add an image centered in the document."""
    if not os.path.exists(img_path):
        p = doc.add_paragraph()
        run = p.add_run(f'[Image not found: {os.path.basename(img_path)}]')
        run.font.color.rgb = RGBColor(0xC1, 0x29, 0x2E)
        run.italic = True
        return

    try:
        with Image.open(img_path) as img:
            w_px, h_px = img.size
            dpi = img.info.get('dpi', (200, 200))
            dpi_x = dpi[0] if isinstance(dpi, tuple) else dpi
            dpi_y = dpi[1] if isinstance(dpi, tuple) else dpi

        w_inches = w_px / dpi_x
        h_inches = h_px / dpi_y

        max_w, max_h = 5.8, 7.0
        if w_inches > max_w:
            scale = max_w / w_inches
            w_inches *= scale
            h_inches *= scale
        if h_inches > max_h:
            scale = max_h / h_inches
            w_inches *= scale
            h_inches *= scale

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run()
        run.add_picture(img_path, width=Inches(w_inches))

    except Exception as e:
        p = doc.add_paragraph()
        run = p.add_run(f'[Error loading image: {e}]')
        run.font.color.rgb = RGBColor(0xC1, 0x29, 0x2E)
        run.italic = True


def process_inline_formatting(paragraph, text):
    """Parse inline Markdown bold/italic and add runs."""
    pattern = re.compile(
        r'(\*\*\*(.+?)\*\*\*)'
        r'|(\*\*(.+?)\*\*)'
        r'|(\*(.+?)\*)'
    )

    last_end = 0
    for match in pattern.finditer(text):
        start = match.start()
        if start > last_end:
            run = paragraph.add_run(text[last_end:start])
            run.font.name = FONT_FAMILY
            run.font.size = FONT_BODY_SIZE
            run.font.color.rgb = TEXT_COLOR

        if match.group(2):
            run = paragraph.add_run(match.group(2))
            run.bold = True
            run.italic = True
        elif match.group(4):
            run = paragraph.add_run(match.group(4))
            run.bold = True
        elif match.group(6):
            run = paragraph.add_run(match.group(6))
            run.italic = True

        run.font.name = FONT_FAMILY
        run.font.size = FONT_BODY_SIZE
        run.font.color.rgb = TEXT_COLOR
        last_end = match.end()

    if last_end < len(text):
        run = paragraph.add_run(text[last_end:])
        run.font.name = FONT_FAMILY
        run.font.size = FONT_BODY_SIZE
        run.font.color.rgb = TEXT_COLOR


def _apply_cell_formatting(cell, text, is_header=False, font_size=Pt(9)):
    """Apply inline formatting to a table cell."""
    cell.text = ''  # Clear default paragraph
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

    # Strip markdown bold/italic markers and detect formatting
    clean = text.strip()
    bold = is_header
    italic = False

    if clean.startswith('**') and clean.endswith('**'):
        clean = clean[2:-2]
        bold = True
    elif clean.startswith('*') and clean.endswith('*'):
        clean = clean[1:-1]
        italic = True

    run = p.add_run(clean)
    run.font.name = FONT_FAMILY
    run.font.size = font_size
    run.font.color.rgb = TEXT_COLOR
    run.bold = bold
    run.italic = italic


def _parse_table_row(line):
    """Parse a Markdown table row into a list of cell strings."""
    stripped = line.strip()
    if stripped.startswith('|'):
        stripped = stripped[1:]
    if stripped.endswith('|'):
        stripped = stripped[:-1]
    return [cell.strip() for cell in stripped.split('|')]


def _is_separator_row(line):
    """Check if a line is a Markdown table separator (|---|---|)."""
    stripped = line.strip().replace(' ', '')
    return bool(re.match(r'^\|?[-:|]+(\|[-:|]+)+\|?$', stripped))


def add_markdown_table(doc, table_lines):
    """Convert Markdown table lines into a styled Word table."""
    if len(table_lines) < 2:
        return

    # Parse header row
    header_cells = _parse_table_row(table_lines[0])
    num_cols = len(header_cells)

    # Find data rows (skip separator)
    data_rows = []
    for line in table_lines[1:]:
        if _is_separator_row(line):
            continue
        cells = _parse_table_row(line)
        # Pad or trim to match column count
        while len(cells) < num_cols:
            cells.append('')
        cells = cells[:num_cols]
        data_rows.append(cells)

    total_rows = 1 + len(data_rows)
    if total_rows < 2:
        return

    # Create Word table
    table = doc.add_table(rows=total_rows, cols=num_cols)
    table.style = 'Table Grid'
    table.autofit = True

    # Header row
    header_row = table.rows[0]
    for j, cell_text in enumerate(header_cells):
        cell = header_row.cells[j]
        _apply_cell_formatting(cell, cell_text, is_header=True, font_size=Pt(9))
        # Header background
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear',
            qn('w:color'): 'auto',
            qn('w:fill'): '1B2A4A',
        })
        shading.append(shd)
        # White text for header
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for i, row_cells in enumerate(data_rows):
        row = table.rows[i + 1]
        for j, cell_text in enumerate(row_cells):
            cell = row.cells[j]
            _apply_cell_formatting(cell, cell_text, is_header=False, font_size=Pt(9))
            # Alternate row shading
            if i % 2 == 1:
                shading = cell._element.get_or_add_tcPr()
                shd = shading.makeelement(qn('w:shd'), {
                    qn('w:val'): 'clear',
                    qn('w:color'): 'auto',
                    qn('w:fill'): 'F0F4F8',
                })
                shading.append(shd)

    # Add spacing after table
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)


def process_guide(doc, guide_entry):
    """Parse a tool guide Markdown file and add it to the Word document."""
    filepath = os.path.join(GUIDES_DIR, guide_entry['file'])
    if not os.path.exists(filepath):
        print(f"  SKIP: {guide_entry['file']} not found")
        return False

    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # Track whether we've inserted the figure (insert after "Step-by-Step" or
    # "Worked Example" or "Interpretation" section heading)
    figure_inserted = False
    figure_insert_headings = [
        'worked example', 'step-by-step', 'interpretation',
        'how to read', 'plotting the result', 'constructing the map',
    ]

    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')
        stripped = line.strip()

        # Skip HTML comments
        if stripped.startswith('<!--'):
            i += 1
            continue

        # Skip horizontal rules
        if stripped == '---':
            i += 1
            continue

        # Empty line
        if stripped == '':
            i += 1
            continue

        # Headings
        if stripped.startswith('#'):
            heading_match = re.match(r'^(#{1,4})\s+(.+)$', stripped)
            if heading_match:
                level = len(heading_match.group(1))
                heading_text = heading_match.group(2).strip()
                heading_clean = re.sub(r'\*+', '', heading_text)
                doc.add_heading(heading_clean, level=min(level, 4))

                # Check if we should insert the figure after this heading's first paragraph
                if not figure_inserted:
                    heading_lower = heading_clean.lower()
                    for trigger in figure_insert_headings:
                        if trigger in heading_lower:
                            # Skip to end of first paragraph after heading, then insert figure
                            i += 1
                            # Skip blanks
                            while i < len(lines) and lines[i].strip() == '':
                                i += 1
                            # Read first paragraph
                            para_lines = []
                            while i < len(lines):
                                nl = lines[i].rstrip('\n').strip()
                                if nl == '' or nl.startswith('#') or nl.startswith('- ') or nl.startswith('* '):
                                    break
                                para_lines.append(nl)
                                i += 1

                            if para_lines:
                                p = doc.add_paragraph()
                                process_inline_formatting(p, ' '.join(para_lines))

                            # Now insert the figure
                            _insert_guide_figure(doc, guide_entry)
                            figure_inserted = True
                            continue

                i += 1
                continue

        # Bullet list
        if stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            process_inline_formatting(p, text)
            i += 1
            continue

        # Numbered list
        num_match = re.match(r'^(\d+)\.\s+(.+)$', stripped)
        if num_match:
            text = num_match.group(2)
            p = doc.add_paragraph(style='List Number')
            process_inline_formatting(p, text)
            i += 1
            continue

        # Markdown table (pipe-delimited rows)
        if stripped.startswith('|'):
            table_lines = []
            j = i
            while j < len(lines):
                tl = lines[j].rstrip('\n').strip()
                if tl.startswith('|'):
                    table_lines.append(tl)
                    j += 1
                else:
                    break
            if len(table_lines) >= 2:
                add_markdown_table(doc, table_lines)
            i = j
            continue

        # Regular paragraph
        para_lines = [stripped]
        j = i + 1
        while j < len(lines):
            next_stripped = lines[j].rstrip('\n').strip()
            if (next_stripped == '' or
                next_stripped.startswith('#') or
                next_stripped.startswith('- ') or
                next_stripped.startswith('* ') or
                next_stripped.startswith('|') or
                re.match(r'^\d+\.\s', next_stripped) or
                next_stripped == '---' or
                next_stripped.startswith('<!--')):
                break
            para_lines.append(next_stripped)
            j += 1

        full_text = ' '.join(para_lines)
        p = doc.add_paragraph()
        process_inline_formatting(p, full_text)
        i = j
        continue

    # If figure wasn't inserted by a trigger heading, insert it at the end
    # (before any Key Terms / References sections if we can detect them)
    if not figure_inserted:
        _insert_guide_figure(doc, guide_entry)

    return True


def _insert_guide_figure(doc, guide_entry):
    """Insert the tool guide's worked-example figure."""
    fig_path = os.path.join(GRAPHICS_DIR, guide_entry['figure'])

    # Figure label
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f'Figure {guide_entry["figure_number"]}. ')
    run.bold = True
    run.font.name = FONT_FAMILY
    run.font.size = Pt(10)
    run.font.color.rgb = TEXT_COLOR
    run = p.add_run(guide_entry['figure_title'])
    run.italic = True
    run.font.name = FONT_FAMILY
    run.font.size = Pt(10)
    run.font.color.rgb = TEXT_SECONDARY

    # Image
    add_image_to_doc(doc, fig_path, guide_entry['alt_text'])


def generate_pdf(docx_path, pdf_path):
    """Convert Word to PDF via PowerShell/Word COM."""
    abs_docx = os.path.abspath(docx_path).replace('\\', '\\\\')
    abs_pdf = os.path.abspath(pdf_path).replace('\\', '\\\\')

    ps_script = f'''
$word = New-Object -ComObject Word.Application
$word.Visible = $false
$doc = $word.Documents.Open("{abs_docx}")
$doc.SaveAs([ref]"{abs_pdf}", [ref]17)
$doc.Close()
$word.Quit()
[System.Runtime.Interopservices.Marshal]::ReleaseComObject($word) | Out-Null
'''

    try:
        result = subprocess.run(
            ['powershell', '-NoProfile', '-Command', ps_script],
            capture_output=True, text=True, timeout=120
        )
        if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
            size_mb = os.path.getsize(pdf_path) / (1024 * 1024)
            print(f"  PDF saved: {pdf_path} ({size_mb:.1f} MB)")
            return True
        else:
            print(f"  PDF generation failed. Error: {result.stderr[:200] if result.stderr else 'unknown'}")
            print(f"  You can open the .docx in Word and Save As PDF manually.")
            return False
    except Exception as e:
        print(f"  PDF generation failed: {e}")
        print(f"  You can open the .docx in Word and Save As PDF manually.")
        return False


# ── Main ────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description='Export tool guides to Word/PDF')
    parser.add_argument('--word-only', action='store_true')
    parser.add_argument('--pdf-only', action='store_true')
    args = parser.parse_args()

    print(f"\n{'='*60}")
    print(f"  TOOL GUIDE EXPORT")
    print(f"  Source: output/tool-guides/")
    print(f"{'='*60}\n")

    if not args.pdf_only:
        doc = create_document()
        add_title_page(doc)
        add_toc(doc)

        guides_processed = 0
        for entry in GUIDE_ORDER:
            print(f"  Processing: {entry['file']}")
            if process_guide(doc, entry):
                guides_processed += 1
                print(f"    -> OK (figure: {entry['figure_number']})")

            # Page break between guides
            doc.add_page_break()

        print(f"\n  {guides_processed} tool guides compiled")
        doc.save(OUTPUT_DOCX)
        size_mb = os.path.getsize(OUTPUT_DOCX) / (1024 * 1024)
        print(f"  Word saved: {OUTPUT_DOCX} ({size_mb:.1f} MB)")

    if not args.word_only:
        if not os.path.exists(OUTPUT_DOCX):
            print("  ERROR: Word file must exist for PDF conversion")
            sys.exit(1)
        print(f"\n  Converting to PDF...")
        generate_pdf(OUTPUT_DOCX, OUTPUT_PDF)

    print(f"\n{'='*60}")
    print(f"  EXPORT COMPLETE")
    print(f"  Word: {OUTPUT_DOCX}")
    if not args.word_only:
        print(f"  PDF:  {OUTPUT_PDF}")
    print(f"{'='*60}\n")


if __name__ == '__main__':
    main()

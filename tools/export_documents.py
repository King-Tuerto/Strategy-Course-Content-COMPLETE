"""
Document Export Tool
=====================
Compiles all enhanced chapters (with figures) into a single Word document
and a single PDF file.

Usage:
    python tools/export_documents.py              # Generate both Word and PDF
    python tools/export_documents.py --word-only   # Word only
    python tools/export_documents.py --pdf-only    # PDF only (requires Word installed)

Output:
    output/Strategy-Course-Complete.docx
    output/Strategy-Course-Complete.pdf
"""

import os
import re
import sys
import argparse

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if PROJECT_ROOT not in sys.path:
    sys.path.insert(0, PROJECT_ROOT)

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from PIL import Image

# ── Paths ───────────────────────────────────────────────────────────────

CHAPTERS_DIR = os.path.join(PROJECT_ROOT, 'output', 'chapters-with-figures')
GRAPHICS_DIR = os.path.join(PROJECT_ROOT, 'output', 'graphics')
OUTPUT_DOCX = os.path.join(PROJECT_ROOT, 'output', 'Strategy-Course-Complete.docx')
OUTPUT_PDF = os.path.join(PROJECT_ROOT, 'output', 'Strategy-Course-Complete.pdf')

# Chapter ordering
CHAPTER_FILES = [
    'Topic-1-Foundations-of-Strategic-Management.md',
    'Topic-2-External-Analysis-and-International-Strategy.md',
    'Topic-3-Internal-Analysis-and-Strategy-Types.md',
    'Topic-4-Strategy-Analysis-and-Implementation.md',
    'Topic-6-Strategy-Evaluation-and-Control.md',
    'Topic-7-Finance-and-Accounting-in-Strategy-Implementation.md',
]

# ── Design Constants ────────────────────────────────────────────────────

NAVY = RGBColor(0x1B, 0x2A, 0x4A)
STEEL = RGBColor(0x3D, 0x5A, 0x80)
TEXT_COLOR = RGBColor(0x21, 0x25, 0x29)
TEXT_SECONDARY = RGBColor(0x49, 0x50, 0x57)
FONT_FAMILY = 'Segoe UI'
FONT_BODY_SIZE = Pt(11)
FONT_H1_SIZE = Pt(22)
FONT_H2_SIZE = Pt(16)
FONT_H3_SIZE = Pt(13)
FONT_H4_SIZE = Pt(11.5)
IMAGE_MAX_WIDTH = Inches(5.8)


# ── Document Builder ────────────────────────────────────────────────────

def create_document():
    """Create a styled Word document."""
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # Default paragraph font
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_FAMILY
    font.size = FONT_BODY_SIZE
    font.color.rgb = TEXT_COLOR
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15

    # Heading styles
    for level, (size, color, bold) in enumerate([
        (FONT_H1_SIZE, NAVY, True),
        (FONT_H2_SIZE, NAVY, True),
        (FONT_H3_SIZE, STEEL, True),
        (FONT_H4_SIZE, STEEL, True),
    ], start=1):
        hstyle = doc.styles[f'Heading {level}']
        hfont = hstyle.font
        hfont.name = FONT_FAMILY
        hfont.size = size
        hfont.color.rgb = color
        hfont.bold = bold
        hpf = hstyle.paragraph_format
        hpf.space_before = Pt(18 if level <= 2 else 12)
        hpf.space_after = Pt(6)
        hpf.keep_with_next = True

    return doc


def add_title_page(doc):
    """Add a styled title page."""
    # Blank spacing
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Strategic Management')
    run.font.size = Pt(36)
    run.font.color.rgb = NAVY
    run.font.name = FONT_FAMILY
    run.bold = True

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Concepts and Applications')
    run.font.size = Pt(20)
    run.font.color.rgb = STEEL
    run.font.name = FONT_FAMILY
    run.italic = True

    # Spacing
    for _ in range(3):
        doc.add_paragraph()

    # Course info
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('MBA Strategy Course')
    run.font.size = Pt(14)
    run.font.color.rgb = TEXT_SECONDARY
    run.font.name = FONT_FAMILY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Grand Canyon University')
    run.font.size = Pt(14)
    run.font.color.rgb = TEXT_SECONDARY
    run.font.name = FONT_FAMILY

    # Page break
    doc.add_page_break()


def add_toc_placeholder(doc):
    """Add a Table of Contents placeholder."""
    p = doc.add_heading('Table of Contents', level=1)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\n[Table of Contents will auto-generate in Word:\n'
                     'Right-click here > Update Field > Update Entire Table]\n')
    run.font.size = Pt(10)
    run.font.color.rgb = TEXT_SECONDARY
    run.font.name = FONT_FAMILY
    run.italic = True

    # Insert TOC field code
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar1 = run._r.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    run._r.append(fldChar1)

    run2 = paragraph.add_run()
    instrText = run2._r.makeelement(qn('w:instrText'), {})
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run2._r.append(instrText)

    run3 = paragraph.add_run()
    fldChar2 = run3._r.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'separate'})
    run3._r.append(fldChar2)

    run4 = paragraph.add_run()
    run4._r.text = 'Press F9 to update table of contents'
    run4.font.color.rgb = TEXT_SECONDARY

    run5 = paragraph.add_run()
    fldChar3 = run5._r.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    run5._r.append(fldChar3)

    doc.add_page_break()


def resolve_image_path(md_path, chapter_file):
    """Resolve a relative image path from markdown to absolute path."""
    # md_path looks like ../graphics/topic-X/filename.png
    # chapter_file is in output/chapters-with-figures/
    chapter_dir = os.path.dirname(os.path.join(CHAPTERS_DIR, chapter_file))
    abs_path = os.path.normpath(os.path.join(chapter_dir, md_path))
    return abs_path


def add_image_to_doc(doc, img_path, alt_text=''):
    """Add an image centered in the document, scaled appropriately."""
    if not os.path.exists(img_path):
        p = doc.add_paragraph()
        run = p.add_run(f'[Image not found: {os.path.basename(img_path)}]')
        run.font.color.rgb = RGBColor(0xC1, 0x29, 0x2E)
        run.italic = True
        return

    try:
        with Image.open(img_path) as img:
            w_px, h_px = img.size
            dpi = img.info.get('dpi', (200, 200))
            dpi_x = dpi[0] if isinstance(dpi, tuple) else dpi

        # Calculate width to fit page
        w_inches = w_px / dpi_x
        h_inches = h_px / (dpi[1] if isinstance(dpi, tuple) else dpi)

        # Scale to fit
        max_w = 5.8  # inches
        max_h = 7.0  # inches
        if w_inches > max_w:
            scale = max_w / w_inches
            w_inches *= scale
            h_inches *= scale
        if h_inches > max_h:
            scale = max_h / h_inches
            w_inches *= scale
            h_inches *= scale

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run()
        run.add_picture(img_path, width=Inches(w_inches))

    except Exception as e:
        p = doc.add_paragraph()
        run = p.add_run(f'[Error loading image: {e}]')
        run.font.color.rgb = RGBColor(0xC1, 0x29, 0x2E)
        run.italic = True


def process_inline_formatting(paragraph, text):
    """
    Parse inline Markdown formatting and add runs to a paragraph.
    Handles **bold**, *italic*, and ***bold italic***.
    """
    # Pattern for bold+italic, bold, or italic
    pattern = re.compile(
        r'(\*\*\*(.+?)\*\*\*)'   # ***bold italic***
        r'|(\*\*(.+?)\*\*)'       # **bold**
        r'|(\*(.+?)\*)'           # *italic*
    )

    last_end = 0
    for match in pattern.finditer(text):
        # Add plain text before this match
        start = match.start()
        if start > last_end:
            run = paragraph.add_run(text[last_end:start])
            run.font.name = FONT_FAMILY
            run.font.size = FONT_BODY_SIZE
            run.font.color.rgb = TEXT_COLOR

        if match.group(2):  # ***bold italic***
            run = paragraph.add_run(match.group(2))
            run.bold = True
            run.italic = True
        elif match.group(4):  # **bold**
            run = paragraph.add_run(match.group(4))
            run.bold = True
        elif match.group(6):  # *italic*
            run = paragraph.add_run(match.group(6))
            run.italic = True

        run.font.name = FONT_FAMILY
        run.font.size = FONT_BODY_SIZE
        run.font.color.rgb = TEXT_COLOR
        last_end = match.end()

    # Remaining text
    if last_end < len(text):
        run = paragraph.add_run(text[last_end:])
        run.font.name = FONT_FAMILY
        run.font.size = FONT_BODY_SIZE
        run.font.color.rgb = TEXT_COLOR


def _apply_cell_formatting(cell, text, is_header=False, font_size=Pt(9)):
    """Apply inline formatting to a table cell."""
    cell.text = ''  # Clear default paragraph
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

    # Strip markdown bold/italic markers and detect formatting
    clean = text.strip()
    bold = is_header
    italic = False

    if clean.startswith('**') and clean.endswith('**'):
        clean = clean[2:-2]
        bold = True
    elif clean.startswith('*') and clean.endswith('*'):
        clean = clean[1:-1]
        italic = True

    run = p.add_run(clean)
    run.font.name = FONT_FAMILY
    run.font.size = font_size
    run.font.color.rgb = TEXT_COLOR
    run.bold = bold
    run.italic = italic


def _parse_table_row(line):
    """Parse a Markdown table row into a list of cell strings."""
    stripped = line.strip()
    if stripped.startswith('|'):
        stripped = stripped[1:]
    if stripped.endswith('|'):
        stripped = stripped[:-1]
    return [cell.strip() for cell in stripped.split('|')]


def _is_separator_row(line):
    """Check if a line is a Markdown table separator (|---|---|)."""
    stripped = line.strip().replace(' ', '')
    return bool(re.match(r'^\|?[-:|]+(\|[-:|]+)+\|?$', stripped))


def add_markdown_table(doc, table_lines):
    """Convert Markdown table lines into a styled Word table."""
    if len(table_lines) < 2:
        return

    # Parse header row
    header_cells = _parse_table_row(table_lines[0])
    num_cols = len(header_cells)

    # Find data rows (skip separator)
    data_rows = []
    for line in table_lines[1:]:
        if _is_separator_row(line):
            continue
        cells = _parse_table_row(line)
        # Pad or trim to match column count
        while len(cells) < num_cols:
            cells.append('')
        cells = cells[:num_cols]
        data_rows.append(cells)

    total_rows = 1 + len(data_rows)
    if total_rows < 2:
        return

    # Create Word table
    table = doc.add_table(rows=total_rows, cols=num_cols)
    table.style = 'Table Grid'
    table.autofit = True

    # Header row
    header_row = table.rows[0]
    for j, cell_text in enumerate(header_cells):
        cell = header_row.cells[j]
        _apply_cell_formatting(cell, cell_text, is_header=True, font_size=Pt(9))
        # Header background
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear',
            qn('w:color'): 'auto',
            qn('w:fill'): '1B2A4A',
        })
        shading.append(shd)
        # White text for header
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for i, row_cells in enumerate(data_rows):
        row = table.rows[i + 1]
        for j, cell_text in enumerate(row_cells):
            cell = row.cells[j]
            _apply_cell_formatting(cell, cell_text, is_header=False, font_size=Pt(9))
            # Alternate row shading
            if i % 2 == 1:
                shading = cell._element.get_or_add_tcPr()
                shd = shading.makeelement(qn('w:shd'), {
                    qn('w:val'): 'clear',
                    qn('w:color'): 'auto',
                    qn('w:fill'): 'F0F4F8',
                })
                shading.append(shd)

    # Add spacing after table
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)


def process_chapter(doc, chapter_file):
    """Parse a Markdown chapter and add it to the Word document."""
    filepath = os.path.join(CHAPTERS_DIR, chapter_file)
    if not os.path.exists(filepath):
        print(f"  SKIP: {chapter_file} not found")
        return 0

    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    figures_added = 0
    i = 0
    in_list = False
    list_buffer = []

    while i < len(lines):
        line = lines[i].rstrip('\n')
        stripped = line.strip()

        # Skip HTML comments (CWV tags)
        if stripped.startswith('<!--'):
            i += 1
            continue

        # Skip horizontal rules
        if stripped == '---':
            i += 1
            continue

        # Empty line
        if stripped == '':
            if in_list and list_buffer:
                in_list = False
                list_buffer = []
            i += 1
            continue

        # Headings
        if stripped.startswith('#'):
            heading_match = re.match(r'^(#{1,4})\s+(.+)$', stripped)
            if heading_match:
                level = len(heading_match.group(1))
                heading_text = heading_match.group(2).strip()
                # Strip markdown formatting from heading
                heading_text = re.sub(r'\*+', '', heading_text)
                doc.add_heading(heading_text, level=min(level, 4))
                i += 1
                continue

        # Figure reference: **Figure X.Y.** *Title*
        fig_match = re.match(r'^\*\*Figure\s+([\d.TG]+)\.\*\*\s+\*(.+?)\*\s*$', stripped)
        if fig_match:
            fig_num = fig_match.group(1)
            fig_title = fig_match.group(2)

            # Add figure label
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(f'Figure {fig_num}. ')
            run.bold = True
            run.font.name = FONT_FAMILY
            run.font.size = Pt(10)
            run.font.color.rgb = TEXT_COLOR
            run = p.add_run(fig_title)
            run.italic = True
            run.font.name = FONT_FAMILY
            run.font.size = Pt(10)
            run.font.color.rgb = TEXT_SECONDARY

            # Look for the image on the next non-blank line
            j = i + 1
            while j < len(lines) and lines[j].strip() == '':
                j += 1

            if j < len(lines):
                img_match = re.match(r'!\[(.+?)\]\((.+?)\)', lines[j].strip())
                if img_match:
                    alt_text = img_match.group(1)
                    img_rel_path = img_match.group(2)
                    img_abs_path = resolve_image_path(img_rel_path, chapter_file)
                    add_image_to_doc(doc, img_abs_path, alt_text)
                    figures_added += 1
                    i = j + 1
                    continue

            i += 1
            continue

        # Image line on its own (without figure label above)
        img_match = re.match(r'!\[(.+?)\]\((.+?)\)', stripped)
        if img_match:
            alt_text = img_match.group(1)
            img_rel_path = img_match.group(2)
            img_abs_path = resolve_image_path(img_rel_path, chapter_file)
            add_image_to_doc(doc, img_abs_path, alt_text)
            figures_added += 1
            i += 1
            continue

        # Bullet list items
        if stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            process_inline_formatting(p, text)
            i += 1
            continue

        # Numbered list items
        num_match = re.match(r'^(\d+)\.\s+(.+)$', stripped)
        if num_match:
            text = num_match.group(2)
            p = doc.add_paragraph(style='List Number')
            process_inline_formatting(p, text)
            i += 1
            continue

        # Markdown table (pipe-delimited rows)
        if stripped.startswith('|'):
            table_lines = []
            j = i
            while j < len(lines):
                tl = lines[j].rstrip('\n').strip()
                if tl.startswith('|'):
                    table_lines.append(tl)
                    j += 1
                else:
                    break
            if len(table_lines) >= 2:
                add_markdown_table(doc, table_lines)
            i = j
            continue

        # Regular paragraph — collect continuation lines
        para_lines = [stripped]
        j = i + 1
        while j < len(lines):
            next_line = lines[j].rstrip('\n')
            next_stripped = next_line.strip()
            # Stop on blank line, heading, list, image, figure ref, table, or HR
            if (next_stripped == '' or
                next_stripped.startswith('#') or
                next_stripped.startswith('- ') or
                next_stripped.startswith('* ') or
                next_stripped.startswith('|') or
                re.match(r'^\d+\.\s', next_stripped) or
                next_stripped.startswith('![') or
                next_stripped.startswith('**Figure') or
                next_stripped == '---' or
                next_stripped.startswith('<!--')):
                break
            para_lines.append(next_stripped)
            j += 1

        full_text = ' '.join(para_lines)
        p = doc.add_paragraph()
        process_inline_formatting(p, full_text)
        i = j
        continue

    return figures_added


def generate_word(doc, output_path):
    """Save the Word document."""
    doc.save(output_path)
    size_mb = os.path.getsize(output_path) / (1024 * 1024)
    print(f"  Word saved: {output_path} ({size_mb:.1f} MB)")


def generate_pdf(docx_path, pdf_path):
    """Convert Word to PDF using Windows COM automation."""
    try:
        import comtypes.client
        word = comtypes.client.CreateObject('Word.Application')
        word.Visible = False

        abs_docx = os.path.abspath(docx_path)
        abs_pdf = os.path.abspath(pdf_path)

        doc = word.Documents.Open(abs_docx)
        doc.SaveAs(abs_pdf, FileFormat=17)  # 17 = wdFormatPDF
        doc.Close()
        word.Quit()

        size_mb = os.path.getsize(pdf_path) / (1024 * 1024)
        print(f"  PDF saved: {pdf_path} ({size_mb:.1f} MB)")
        return True

    except ImportError:
        print("  NOTE: comtypes not installed. Attempting alternative PDF method...")
        return _generate_pdf_powershell(docx_path, pdf_path)
    except Exception as e:
        print(f"  NOTE: COM automation failed ({e}). Attempting PowerShell method...")
        return _generate_pdf_powershell(docx_path, pdf_path)


def _generate_pdf_powershell(docx_path, pdf_path):
    """Fallback: use PowerShell to convert Word to PDF."""
    import subprocess
    abs_docx = os.path.abspath(docx_path).replace('\\', '\\\\')
    abs_pdf = os.path.abspath(pdf_path).replace('\\', '\\\\')

    ps_script = f'''
$word = New-Object -ComObject Word.Application
$word.Visible = $false
$doc = $word.Documents.Open("{abs_docx}")
$doc.SaveAs([ref]"{abs_pdf}", [ref]17)
$doc.Close()
$word.Quit()
[System.Runtime.Interopservices.Marshal]::ReleaseComObject($word) | Out-Null
'''

    try:
        result = subprocess.run(
            ['powershell', '-NoProfile', '-Command', ps_script],
            capture_output=True, text=True, timeout=120
        )
        if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
            size_mb = os.path.getsize(pdf_path) / (1024 * 1024)
            print(f"  PDF saved: {pdf_path} ({size_mb:.1f} MB)")
            return True
        else:
            print(f"  PDF generation failed. Error: {result.stderr[:200] if result.stderr else 'unknown'}")
            print(f"  You can open the .docx in Word and Save As PDF manually.")
            return False
    except Exception as e:
        print(f"  PDF generation failed: {e}")
        print(f"  You can open the .docx in Word and Save As PDF manually.")
        return False


# ── Main ────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description='Export course to Word/PDF')
    parser.add_argument('--word-only', action='store_true', help='Only generate Word')
    parser.add_argument('--pdf-only', action='store_true', help='Only generate PDF')
    args = parser.parse_args()

    print(f"\n{'='*60}")
    print(f"  DOCUMENT EXPORT TOOL")
    print(f"  Source: output/chapters-with-figures/")
    print(f"{'='*60}\n")

    # Build Word document
    if not args.pdf_only:
        doc = create_document()
        add_title_page(doc)
        add_toc_placeholder(doc)

        total_figures = 0
        for chapter_file in CHAPTER_FILES:
            print(f"  Processing: {chapter_file}")
            figs = process_chapter(doc, chapter_file)
            total_figures += figs
            print(f"    -> {figs} figures embedded")

            # Page break between chapters
            doc.add_page_break()

        print(f"\n  Total figures embedded: {total_figures}")
        generate_word(doc, OUTPUT_DOCX)

    # Generate PDF
    if not args.word_only:
        if not os.path.exists(OUTPUT_DOCX):
            print("  ERROR: Word file must be generated first for PDF conversion")
            sys.exit(1)

        print(f"\n  Converting to PDF...")
        generate_pdf(OUTPUT_DOCX, OUTPUT_PDF)

    print(f"\n{'='*60}")
    print(f"  EXPORT COMPLETE")
    print(f"  Word: {OUTPUT_DOCX}")
    if not args.word_only:
        print(f"  PDF:  {OUTPUT_PDF}")
    print(f"{'='*60}\n")


if __name__ == '__main__':
    main()
